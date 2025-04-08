import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import time
from tqdm import tqdm
import argparse
import re
from collections import defaultdict
import logging
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from concurrent.futures import ProcessPoolExecutor, as_completed

# Import existing processing modules
from process_pdfs import process_single_pdf, load_lexicon, configure_logging, format_time, create_output_directories
from term_counter import MatchClassifier

def extract_year_from_path(path):
    """Extract year from the directory path."""
    year_match = re.search(r'(\d{4})_StudiesWDocs', path)
    if year_match:
        return int(year_match.group(1))
    return None

def process_pdfs_by_year(base_dir, lexicon_file, output_folder, threshold=85, workers=None, sample_mode=False, sample_size=5):
    """
    Process all PDFs grouped by year and analyze trends over time.
    
    Args:
        base_dir: Directory containing year-based subdirectories with clinical trials
        lexicon_file: CSV file containing lexicon terms to match
        output_folder: Folder for outputting results
        threshold: Matching threshold (default: 85)
        workers: Number of worker processes for parallel processing
        sample_mode: If True, only process a limited number of files per year (for testing)
        sample_size: Number of files to process per year when in sample mode
    """
    # Configure logging
    configure_logging()
    
    start_time = time.time()
    
    # Create timestamped output directory
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    timestamped_output_dir = os.path.join(output_folder, f"temporal_analysis_{timestamp}")
    os.makedirs(timestamped_output_dir, exist_ok=True)
    
    # Create organized subdirectories
    output_dirs = create_output_directories(timestamped_output_dir)
    
    # Load lexicon
    print(f"\nLoading lexicon from {lexicon_file}...")
    lexicon_load_start = time.time()
    lexicon_terms = load_lexicon(lexicon_file)
    print(f"Loaded {len(lexicon_terms)} lexicon terms in {format_time(time.time() - lexicon_load_start)}")
    
    # Find all year directories
    year_dirs = []
    for item in os.listdir(base_dir):
        item_path = os.path.join(base_dir, item)
        if os.path.isdir(item_path) and re.match(r'\d{4}_StudiesWDocs', item):
            year_dirs.append(item_path)
    
    if not year_dirs:
        print(f"No year directories found in {base_dir}")
        return
    
    year_dirs.sort()  # Sort to process in chronological order
    print(f"Found {len(year_dirs)} year directories to process: {[os.path.basename(d) for d in year_dirs]}")
    
    # Store results for each year
    yearly_results = {}
    yearly_metrics = {}
    
    # Process each year directory
    for year_dir in year_dirs:
        year = extract_year_from_path(year_dir)
        if not year:
            continue
            
        print(f"\nProcessing files from {year}...")
        
        # Find all PDF files in this year directory
        pdf_files = []
        for root, _, files in os.walk(year_dir):
            for filename in files:
                if filename.endswith(".pdf"):
                    full_path = os.path.join(root, filename)
                    rel_path = os.path.relpath(full_path, year_dir)
                    trial_id = None
                    
                    # Extract trial ID from path
                    path_parts = full_path.split(os.sep)
                    for part in path_parts:
                        if part.startswith("NCT"):
                            trial_id = part
                            break
                            
                    pdf_files.append((full_path, rel_path, filename, trial_id))
        
        # Limit files in sample mode
        if sample_mode and len(pdf_files) > sample_size:
            print(f"Sample mode: limiting to {sample_size} files per year")
            pdf_files = pdf_files[:sample_size]
        
        total_files = len(pdf_files)
        print(f"Found {total_files} PDF files for {year}")
        
        # Skip if no files found
        if total_files == 0:
            continue
        
        # Process files for this year
        all_results = []
        
        # Create processing tasks
        processing_tasks = [(pdf_path, rel_path, filename, lexicon_terms, threshold) 
                            for pdf_path, rel_path, filename, _ in pdf_files]
        
        # Initialize progress bar
        progress_bar = tqdm(total=total_files, desc=f"Processing {year} PDFs", unit="file")
        
        # Use ProcessPoolExecutor for parallel processing
        with ProcessPoolExecutor(max_workers=workers) as executor:
            futures = {executor.submit(process_single_pdf, task): task for task in processing_tasks}
            
            for future in as_completed(futures):
                try:
                    result = future.result()
                    if result["has_matches"]:
                        all_results.extend(result["matches"])
                except Exception as e:
                    pdf_path = futures[future][1] if future in futures else "unknown"
                    logging.error(f"Failed to process file {pdf_path}: {str(e)}")
                
                # Update progress bar
                progress_bar.update(1)
        
        # Close progress bar
        progress_bar.close()
        
        if all_results:
            # Add year to each result
            for result in all_results:
                result['year'] = year
                
            # Save results for this year
            yearly_results[year] = all_results
            
            # Create DataFrame for this year's results
            year_df = pd.DataFrame(all_results)
            
            # Save to CSV
            csv_output_path = os.path.join(output_dirs['csv'], f"term_locations_{year}.csv")
            year_df.to_csv(csv_output_path, index=False)
            
            # Generate metrics for this year
            classifier = MatchClassifier(all_results)
            
            # Calculate metrics
            files_with_innovation = len(year_df['rel_path'].unique())
            innovation_percentage = round((files_with_innovation / total_files * 100), 2)
            
            # Category statistics
            category_stats = year_df.groupby('category').agg({
                'rel_path': 'nunique',
                'matched_term': 'count'
            }).rename(columns={
                'rel_path': 'unique_files',
                'matched_term': 'total_matches'
            })
            
            # Calculate percentages
            category_stats['percentage_of_files'] = round((category_stats['unique_files'] / total_files * 100), 2)
            category_stats['matches_per_file'] = round((category_stats['total_matches'] / category_stats['unique_files']), 2)
            
            # Store yearly metrics
            yearly_metrics[year] = {
                'total_files': total_files,
                'files_with_innovation': files_with_innovation,
                'innovation_percentage': innovation_percentage,
                'category_stats': category_stats,
                'classifier': classifier
            }
            
            print(f"Processed {year}: {files_with_innovation} out of {total_files} files ({innovation_percentage}%) contain innovations")
        else:
            print(f"No innovation matches found in {year} files")
    
    # Create temporal analysis after processing all years
    if yearly_metrics:
        create_temporal_analysis(yearly_metrics, lexicon_terms, output_dirs, timestamp)
    else:
        print("No data to generate temporal analysis")
    
    total_duration = time.time() - start_time
    print(f"\nTotal processing time: {format_time(total_duration)}")
    print(f"\nTemporal analysis results saved to: {timestamped_output_dir}")

def create_temporal_analysis(yearly_metrics, lexicon_terms, output_dirs, timestamp):
    """
    Create temporal analysis visualizations and report.
    
    Args:
        yearly_metrics: Dictionary of metrics for each year
        lexicon_terms: List of lexicon terms
        output_dirs: Dictionary of output directories
        timestamp: Timestamp string
    """
    print("\nGenerating temporal analysis...")
    
    # Set up plot style
    sns.set_style("whitegrid")
    plt.rcParams['figure.facecolor'] = 'white'
    
    # Track which visualizations were successfully created
    successful_charts = []
    
    # Try each visualization separately and continue even if some fail
    try:
        # 1. Overall innovation trends over time
        create_overall_innovation_trend(yearly_metrics, output_dirs)
        successful_charts.append('overall_innovation_trend')
    except Exception as e:
        print(f"Error creating overall innovation trend chart: {str(e)}")
    
    try:
        # 2. Category prevalence over time
        create_category_prevalence_trend(yearly_metrics, output_dirs)
        successful_charts.append('category_prevalence')
    except Exception as e:
        print(f"Error creating category prevalence trend chart: {str(e)}")
    
    try:
        # 3. Innovation complexity over time
        create_innovation_complexity_trend(yearly_metrics, output_dirs)
        successful_charts.append('innovation_complexity')
    except Exception as e:
        print(f"Error creating innovation complexity trend chart: {str(e)}")
    
    try:
        # 4. Category growth rates (if we have multiple years)
        if len(yearly_metrics.keys()) >= 2:
            create_category_growth_chart(yearly_metrics, output_dirs)
            successful_charts.append('category_growth')
    except Exception as e:
        print(f"Error creating category growth chart: {str(e)}")
    
    try:
        # 5. Term type distribution over time
        create_term_type_trend(yearly_metrics, output_dirs)
        successful_charts.append('term_type')
    except Exception as e:
        print(f"Error creating term type trend chart: {str(e)}")
    
    try:
        # 6. Generate temporal analysis CSV
        generate_temporal_data_csv(yearly_metrics, output_dirs)
    except Exception as e:
        print(f"Error generating temporal data CSV files: {str(e)}")
    
    try:
        # 7. Create Word document report with available visualizations
        create_temporal_report(yearly_metrics, lexicon_terms, output_dirs, timestamp, successful_charts)
    except Exception as e:
        print(f"Error creating temporal report: {str(e)}")
        
    print(f"Successfully created {len(successful_charts)} visualizations: {successful_charts}")

def create_overall_innovation_trend(yearly_metrics, output_dirs):
    """Create chart showing overall innovation adoption over time."""
    years = sorted(yearly_metrics.keys())
    innovation_percentages = [yearly_metrics[year]['innovation_percentage'] for year in years]
    
    plt.figure(figsize=(10, 6))
    plt.plot(years, innovation_percentages, marker='o', linestyle='-', linewidth=2)
    
    # Add labels and title
    plt.title('Clinical Trial Innovation Adoption Over Time', fontsize=16)
    plt.xlabel('Year', fontsize=14)
    plt.ylabel('Percentage of Files with Innovations (%)', fontsize=14)
    plt.xticks(years)
    plt.grid(True)
    
    # Add values above points
    for i, v in enumerate(innovation_percentages):
        plt.text(years[i], v + 1, f"{v}%", ha='center')
    
    # Add trendline
    if len(years) > 1:
        z = np.polyfit(years, innovation_percentages, 1)
        p = np.poly1d(z)
        plt.plot(years, p(years), "r--", alpha=0.8, label=f"Trend line (slope: {z[0]:.2f})")
        plt.legend()
    
    plt.tight_layout()
    plt.savefig(os.path.join(output_dirs['png'], 'overall_innovation_trend.png'))
    plt.close()

def create_category_prevalence_trend(yearly_metrics, output_dirs):
    """Create chart showing category prevalence trends over time."""
    years = sorted(yearly_metrics.keys())
    
    # Collect all categories across all years
    all_categories = set()
    for year in years:
        all_categories.update(yearly_metrics[year]['category_stats'].index)
    
    # Create DataFrame for visualization
    data = []
    for year in years:
        stats = yearly_metrics[year]['category_stats']
        for category in all_categories:
            if category in stats.index:
                percentage = stats.loc[category, 'percentage_of_files']
            else:
                percentage = 0
            data.append({
                'Year': year,
                'Category': category,
                'Percentage': percentage
            })
    
    df = pd.DataFrame(data)
    
    # Plot top categories (to avoid overcrowding)
    top_categories = df.groupby('Category')['Percentage'].mean().nlargest(10).index
    
    # Filter for top categories
    df_top = df[df['Category'].isin(top_categories)]
    
    plt.figure(figsize=(12, 8))
    
    # Create line plot for each category
    sns.lineplot(data=df_top, x='Year', y='Percentage', hue='Category', marker='o', linewidth=2)
    
    plt.title('Top Innovation Categories Prevalence Over Time', fontsize=16)
    plt.xlabel('Year', fontsize=14)
    plt.ylabel('Percentage of Files (%)', fontsize=14)
    plt.xticks(years)
    plt.legend(title='Category', bbox_to_anchor=(1.05, 1), loc='upper left')
    
    plt.tight_layout()
    plt.savefig(os.path.join(output_dirs['png'], 'category_prevalence_trend.png'))
    plt.close()
    
    # Create heatmap for all categories
    pivot_df = df.pivot(index='Category', columns='Year', values='Percentage')
    
    plt.figure(figsize=(12, 10))
    sns.heatmap(pivot_df, annot=True, fmt='.1f', cmap='YlGnBu', linewidths=.5)
    
    plt.title('Innovation Categories Prevalence Heatmap (%)', fontsize=16)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dirs['png'], 'category_prevalence_heatmap.png'))
    plt.close()

def create_innovation_complexity_trend(yearly_metrics, output_dirs):
    """Create chart showing innovation complexity (categories per file) over time."""
    years = sorted(yearly_metrics.keys())
    avg_categories = []
    
    for year in years:
        # Get all files with innovation for this year
        classifier = yearly_metrics[year]['classifier']
        term_relationships = classifier.get_term_relationships()
        
        # Calculate average categories per file
        if term_relationships:
            categories_per_file = [len(cats) for cats in term_relationships.values()]
            avg = sum(categories_per_file) / len(categories_per_file)
        else:
            avg = 0
            
        avg_categories.append(avg)
    
    plt.figure(figsize=(10, 6))
    plt.plot(years, avg_categories, marker='o', linestyle='-', linewidth=2)
    
    # Add labels and title
    plt.title('Innovation Complexity Over Time', fontsize=16)
    plt.xlabel('Year', fontsize=14)
    plt.ylabel('Average Categories per File', fontsize=14)
    plt.xticks(years)
    plt.grid(True)
    
    # Add values above points
    for i, v in enumerate(avg_categories):
        plt.text(years[i], v + 0.05, f"{v:.2f}", ha='center')
    
    # Add trendline
    if len(years) > 1:
        z = np.polyfit(years, avg_categories, 1)
        p = np.poly1d(z)
        plt.plot(years, p(years), "r--", alpha=0.8, label=f"Trend line (slope: {z[0]:.2f})")
        plt.legend()
    
    plt.tight_layout()
    plt.savefig(os.path.join(output_dirs['png'], 'innovation_complexity_trend.png'))
    plt.close()

def create_category_growth_chart(yearly_metrics, output_dirs):
    """Create chart showing category growth rates between years."""
    years = sorted(yearly_metrics.keys())
    
    if len(years) < 2:
        print("Need at least two years to calculate growth rates")
        return
    
    # Collect all categories
    all_categories = set()
    for year in years:
        all_categories.update(yearly_metrics[year]['category_stats'].index)
    
    # Calculate growth rates
    growth_data = []
    
    for i in range(1, len(years)):
        prev_year = years[i-1]
        curr_year = years[i]
        
        prev_stats = yearly_metrics[prev_year]['category_stats']
        curr_stats = yearly_metrics[curr_year]['category_stats']
        
        for category in all_categories:
            # Safely access values - check if category exists in both years
            if category in prev_stats.index and category in curr_stats.index:
                prev_pct = prev_stats.loc[category, 'percentage_of_files']
                curr_pct = curr_stats.loc[category, 'percentage_of_files']
                
                # Calculate growth (percentage points change)
                growth = curr_pct - prev_pct
                
                growth_data.append({
                    'Period': f"{prev_year}-{curr_year}",
                    'Category': category,
                    'Growth': growth
                })
    
    # If no growth data found (no categories appear in consecutive years), return
    if not growth_data:
        print("No valid growth data found between consecutive years")
        return
        
    growth_df = pd.DataFrame(growth_data)
    
    # Get top growing and declining categories if we have enough data
    if len(growth_df['Category'].unique()) >= 2:
        avg_growth = growth_df.groupby('Category')['Growth'].mean()
        
        # Limit to available categories (may have fewer than 5)
        top_count = min(5, len(avg_growth) // 2) if len(avg_growth) > 1 else 1
        
        if top_count > 0:
            top_growing = avg_growth.nlargest(top_count).index.tolist()
            top_declining = avg_growth.nsmallest(top_count).index.tolist()
            
            # Filter for these categories
            top_categories = top_growing + top_declining
            top_growth_df = growth_df[growth_df['Category'].isin(top_categories)]
            
            # Create plot
            plt.figure(figsize=(12, 8))
            
            # Create bar plot
            ax = sns.barplot(data=top_growth_df, x='Category', y='Growth', hue='Period')
            
            plt.title('Top Growing & Declining Innovation Categories', fontsize=16)
            plt.xlabel('Category', fontsize=14)
            plt.ylabel('Growth (Percentage Points)', fontsize=14)
            plt.xticks(rotation=45, ha='right')
            plt.axhline(y=0, color='black', linestyle='-', alpha=0.3)
            plt.legend(title='Period')
            
            # Add value labels
            for container in ax.containers:
                ax.bar_label(container, fmt='%.1f')
            
            plt.tight_layout()
            plt.savefig(os.path.join(output_dirs['png'], 'category_growth_chart.png'))
            plt.close()
        else:
            print("Not enough distinct categories with growth data to create chart")
    else:
        print("Not enough categories to calculate meaningful growth rates")

def create_term_type_trend(yearly_metrics, output_dirs):
    """Create chart showing primary vs related term usage over time."""
    years = sorted(yearly_metrics.keys())
    
    # Collect primary and related term counts for each year
    primary_counts = []
    related_counts = []
    
    for year in years:
        classifier = yearly_metrics[year]['classifier']
        stats = classifier.get_category_statistics()
        
        primary_count = stats['primary_term_matches'].sum()
        related_count = stats['related_term_matches'].sum()
        
        primary_counts.append(primary_count)
        related_counts.append(related_count)
    
    # Create DataFrame for plotting
    data = []
    for i, year in enumerate(years):
        data.append({'Year': year, 'Type': 'Primary Terms', 'Count': primary_counts[i]})
        data.append({'Year': year, 'Type': 'Related Terms', 'Count': related_counts[i]})
    
    df = pd.DataFrame(data)
    
    # Create line plot
    plt.figure(figsize=(10, 6))
    sns.lineplot(data=df, x='Year', y='Count', hue='Type', marker='o', linewidth=2)
    
    plt.title('Primary vs Related Term Usage Over Time', fontsize=16)
    plt.xlabel('Year', fontsize=14)
    plt.ylabel('Number of Matches', fontsize=14)
    plt.xticks(years)
    plt.legend(title='Term Type')
    
    plt.tight_layout()
    plt.savefig(os.path.join(output_dirs['png'], 'term_type_trend.png'))
    plt.close()
    
    # Create stacked bar chart
    plt.figure(figsize=(10, 6))
    
    # Plot stacked bars
    bar_width = 0.8
    plt.bar(years, primary_counts, label='Primary Terms', width=bar_width)
    plt.bar(years, related_counts, bottom=primary_counts, label='Related Terms', width=bar_width)
    
    # Add labels
    for i in range(len(years)):
        # Add primary term count label
        if primary_counts[i] > 0:
            plt.text(years[i], primary_counts[i]/2, str(primary_counts[i]), 
                     ha='center', va='center', color='white', fontweight='bold')
        
        # Add related term count label
        if related_counts[i] > 0:
            plt.text(years[i], primary_counts[i] + related_counts[i]/2, str(related_counts[i]), 
                     ha='center', va='center', color='white', fontweight='bold')
        
        # Add total label
        total = primary_counts[i] + related_counts[i]
        plt.text(years[i], total + 10, f"Total: {total}", ha='center', va='bottom')
    
    plt.title('Primary vs Related Term Usage by Year', fontsize=16)
    plt.xlabel('Year', fontsize=14)
    plt.ylabel('Number of Matches', fontsize=14)
    plt.xticks(years)
    plt.legend()
    
    plt.tight_layout()
    plt.savefig(os.path.join(output_dirs['png'], 'term_type_stacked.png'))
    plt.close()

def generate_temporal_data_csv(yearly_metrics, output_dirs):
    """Generate CSV files with temporal analysis data."""
    years = sorted(yearly_metrics.keys())
    
    # 1. Overall metrics by year
    overall_data = []
    for year in years:
        metrics = yearly_metrics[year]
        overall_data.append({
            'Year': year,
            'Total_Files': metrics['total_files'],
            'Files_With_Innovation': metrics['files_with_innovation'],
            'Innovation_Percentage': metrics['innovation_percentage']
        })
    
    overall_df = pd.DataFrame(overall_data)
    overall_df.to_csv(os.path.join(output_dirs['csv'], 'temporal_overall_metrics.csv'), index=False)
    
    # 2. Category metrics by year
    category_data = []
    
    for year in years:
        stats = yearly_metrics[year]['category_stats']
        for category, row in stats.iterrows():
            category_data.append({
                'Year': year,
                'Category': category,
                'Unique_Files': row['unique_files'],
                'Total_Matches': row['total_matches'],
                'Percentage_of_Files': row['percentage_of_files'],
                'Matches_Per_File': row['matches_per_file']
            })
    
    category_df = pd.DataFrame(category_data)
    category_df.to_csv(os.path.join(output_dirs['csv'], 'temporal_category_metrics.csv'), index=False)
    
    # 3. Term type statistics by year
    term_type_data = []
    
    for year in years:
        classifier = yearly_metrics[year]['classifier']
        stats = classifier.get_category_statistics()
        
        for _, row in stats.iterrows():
            term_type_data.append({
                'Year': year,
                'Category': row['category'],
                'Primary_Term_Matches': row['primary_term_matches'],
                'Related_Term_Matches': row['related_term_matches'],
                'Most_Common_Terms': ', '.join(list(row['most_common_terms'].keys())[:3])
            })
    
    term_type_df = pd.DataFrame(term_type_data)
    term_type_df.to_csv(os.path.join(output_dirs['csv'], 'temporal_term_type_metrics.csv'), index=False)

def create_temporal_report(yearly_metrics, lexicon_terms, output_dirs, timestamp, successful_charts):
    """Create a Word document report for temporal analysis."""
    years = sorted(yearly_metrics.keys())
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Clinical Trial Innovation Temporal Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    doc.add_paragraph(f'Analysis Date: {timestamp}')
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run('Temporal Analysis Overview:\n').bold = True
    summary.add_run(f"• Analysis covers clinical trials from {min(years)} to {max(years)}\n")
    summary.add_run(f"• {len(lexicon_terms)} innovation categories were analyzed for trends\n")
    
    # Add trend highlight if we have multiple years
    if len(years) > 1:
        # Calculate overall trend
        first_year = min(years)
        last_year = max(years)
        first_pct = yearly_metrics[first_year]['innovation_percentage']
        last_pct = yearly_metrics[last_year]['innovation_percentage']
        change = last_pct - first_pct
        
        if change > 0:
            trend_text = f"• Overall innovation adoption INCREASED by {abs(change):.1f} percentage points from {first_year} to {last_year}\n"
        elif change < 0:
            trend_text = f"• Overall innovation adoption DECREASED by {abs(change):.1f} percentage points from {first_year} to {last_year}\n"
        else:
            trend_text = f"• Overall innovation adoption remained UNCHANGED from {first_year} to {last_year}\n"
            
        summary.add_run(trend_text)
    
    doc.add_paragraph()
    
    # Overall Innovation Trend
    doc.add_heading('Innovation Adoption Over Time', level=1)
    img_path = os.path.join(output_dirs['png'], 'overall_innovation_trend.png')
    if 'overall_innovation_trend' in successful_charts and os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 1: Percentage of Clinical Trial Files with Innovation Terms by Year')
    
    # Innovation by Year Table
    doc.add_heading('Innovation Statistics by Year', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ['Year', 'Total Files', 'Files with Innovation', 'Innovation %']
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    for year in years:
        metrics = yearly_metrics[year]
        row_cells = table.add_row().cells
        row_cells[0].text = str(year)
        row_cells[1].text = str(metrics['total_files'])
        row_cells[2].text = str(metrics['files_with_innovation'])
        row_cells[3].text = f"{metrics['innovation_percentage']}%"
    
    doc.add_paragraph()
    
    # Category Prevalence Trends
    doc.add_heading('Innovation Category Trends', level=1)
    
    img_path = os.path.join(output_dirs['png'], 'category_prevalence_trend.png')
    if 'category_prevalence' in successful_charts and os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 2: Top Innovation Categories Prevalence Over Time')
    
    img_path = os.path.join(output_dirs['png'], 'category_prevalence_heatmap.png')
    if 'category_prevalence' in successful_charts and os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 3: Innovation Categories Prevalence Heatmap')
    
    # Add category growth chart if we have multiple years
    if 'category_growth' in successful_charts and len(years) > 1:
        img_path = os.path.join(output_dirs['png'], 'category_growth_chart.png')
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6))
            doc.add_paragraph('Figure 4: Top Growing & Declining Innovation Categories')
    
    # Innovation Complexity Trend
    doc.add_heading('Innovation Complexity Over Time', level=1)
    img_path = os.path.join(output_dirs['png'], 'innovation_complexity_trend.png')
    if 'innovation_complexity' in successful_charts and os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 5: Average Innovation Categories per File by Year')
    
    # Term Type Analysis
    doc.add_heading('Term Type Analysis Over Time', level=1)
    img_path = os.path.join(output_dirs['png'], 'term_type_stacked.png')
    if 'term_type' in successful_charts and os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 6: Primary vs Related Term Usage by Year')
    
    # Year-by-Year Analysis
    doc.add_heading('Year-by-Year Category Analysis', level=1)
    
    for year in years:
        doc.add_heading(f"{year} Analysis", level=2)
        metrics = yearly_metrics[year]
        p = doc.add_paragraph()
        p.add_run(f'Files with innovation: {metrics["files_with_innovation"]} out of {metrics["total_files"]} ({metrics["innovation_percentage"]}%)\n\n').bold = True
        
        # Add top categories for this year
        p.add_run('Top Categories:\n').bold = True
        stats = metrics['category_stats']
        top_categories = stats.sort_values('percentage_of_files', ascending=False).head(5)
        
        for category, row in top_categories.iterrows():
            p.add_run(f"• {category}: {row['percentage_of_files']}% of files, {row['matches_per_file']} matches per file\n")
        
        # Add most common terms
        p.add_run('\nMost Common Terms:\n').bold = True
        classifier = metrics['classifier']
        term_stats = classifier.get_category_statistics()
        
        for _, row in term_stats.head(5).iterrows():
            common_terms = list(row['most_common_terms'].keys())[:3]
            if common_terms:
                p.add_run(f"• {row['category']}: {', '.join(common_terms)}\n")
    
    # Conclusions and Insights
    if len(years) > 1:
        doc.add_heading('Conclusions and Insights', level=1)
        conclusions = doc.add_paragraph()
        
        # Add overall trend observation
        first_year = min(years)
        last_year = max(years)
        first_pct = yearly_metrics[first_year]['innovation_percentage']
        last_pct = yearly_metrics[last_year]['innovation_percentage']
        change = last_pct - first_pct
        
        if change > 0:
            conclusions.add_run(f"• Innovation adoption has increased by {abs(change):.1f} percentage points from {first_year} to {last_year}, showing positive modernization of clinical trials.\n")
        elif change < 0:
            conclusions.add_run(f"• Innovation adoption has decreased by {abs(change):.1f} percentage points from {first_year} to {last_year}, suggesting potential barriers to modernization.\n")
        else:
            conclusions.add_run(f"• Innovation adoption has remained stable from {first_year} to {last_year} with no significant change.\n")
        
        # Identify fastest growing categories
        if len(years) >= 2:
            category_data = []
            for year in years:
                stats = yearly_metrics[year]['category_stats']
                for category, row in stats.iterrows():
                    category_data.append({
                        'Year': year,
                        'Category': category,
                        'Percentage': row['percentage_of_files']
                    })
            
            category_df = pd.DataFrame(category_data)
            
            # For each category, calculate growth from first to last year
            growth_data = {}
            for category in category_df['Category'].unique():
                cat_data = category_df[category_df['Category'] == category]
                if first_year in cat_data['Year'].values and last_year in cat_data['Year'].values:
                    first_val = cat_data[cat_data['Year'] == first_year]['Percentage'].values[0]
                    last_val = cat_data[cat_data['Year'] == last_year]['Percentage'].values[0]
                    growth_data[category] = last_val - first_val
            
            # Get top growing categories
            if growth_data:
                top_growing = sorted(growth_data.items(), key=lambda x: x[1], reverse=True)[:3]
                
                if top_growing and top_growing[0][1] > 0:
                    conclusions.add_run("\n• Fastest growing innovation categories:\n")
                    for category, growth in top_growing:
                        if growth > 0:
                            conclusions.add_run(f"  - {category}: +{growth:.1f} percentage points\n")
    
    # Generated Files
    doc.add_heading('Generated Analysis Files', level=1)
    files_section = doc.add_paragraph()
    
    # List files by directory
    for dir_type, dir_path in output_dirs.items():
        if dir_type != 'report':  # Skip report directory in listing
            files_section.add_run(f'\n{dir_type.upper()} Files:\n').bold = True
            files = [f for f in os.listdir(dir_path) if os.path.isfile(os.path.join(dir_path, f))]
            for file in sorted(files):
                files_section.add_run(f"• {file}\n")
    
    # Save the document
    doc_path = os.path.join(output_dirs['report'], 'temporal_innovation_analysis.docx')
    doc.save(doc_path)
    return doc_path

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Analyze clinical trial innovation trends over time")
    parser.add_argument("base_dir", help="Base directory containing year-based clinical trial directories")
    parser.add_argument("lexicon_file", help="CSV file containing lexicon terms to match")
    parser.add_argument("output_folder", help="Folder for outputting results")
    parser.add_argument("--threshold", type=int, default=85, help="Matching threshold (default: 85)")
    parser.add_argument("--workers", type=int, help="Number of worker processes for parallel processing")
    parser.add_argument("--verbose", action="store_true", help="Enable verbose output")
    parser.add_argument("--sample", action="store_true", help="Run in sample mode (limit files per year)")
    parser.add_argument("--sample-size", type=int, default=5, help="Number of files to process per year in sample mode")
    
    args = parser.parse_args()
    
    # Set logging level based on verbosity
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
        print("Verbose mode enabled - showing detailed progress information")
    
    process_pdfs_by_year(
        args.base_dir, 
        args.lexicon_file, 
        args.output_folder, 
        args.threshold, 
        args.workers, 
        args.sample, 
        args.sample_size
    ) 