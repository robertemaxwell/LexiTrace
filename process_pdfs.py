import os
import pandas as pd
import time
from datetime import timedelta
import matplotlib.pyplot as plt
import seaborn as sns
from term_counter import load_lexicon, process_pdf_for_terms, MatchClassifier
from pdf_highlighter import highlight_terms_in_pdf
from tqdm import tqdm
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import multiprocessing
import argparse
from concurrent.futures import ProcessPoolExecutor, as_completed

def create_output_directories(base_dir):
    """Create organized subdirectories for different output types."""
    dirs = {
        'png': os.path.join(base_dir, 'visualizations'),
        'csv': os.path.join(base_dir, 'data'),
        'pdf': os.path.join(base_dir, 'highlighted_pdfs'),
        'report': os.path.join(base_dir, 'report')
    }
    for dir_path in dirs.values():
        os.makedirs(dir_path, exist_ok=True)
    return dirs

def format_time(seconds):
    """Format time duration in a human-readable format."""
    return str(timedelta(seconds=round(seconds)))

def create_category_distribution_chart(metrics, output_dir):
    """Create a bar chart showing the distribution of innovation categories."""
    plt.figure(figsize=(12, 6))
    stats = metrics['category_stats']
    
    # Create bar chart
    ax = sns.barplot(x=stats.index, y='percentage_of_files', data=stats)
    
    # Customize the chart
    plt.title('Distribution of Innovation Categories Across Files', pad=20)
    plt.xlabel('Innovation Category')
    plt.ylabel('Percentage of Files (%)')
    plt.xticks(rotation=45, ha='right')
    
    # Add value labels on top of bars
    for i, v in enumerate(stats['percentage_of_files']):
        ax.text(i, v, f'{v}%', ha='center', va='bottom')
    
    # Adjust layout and save
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir['png'], 'category_distribution.png'))
    plt.close()

def create_co_occurrence_heatmap(metrics, output_dir):
    """Create a heatmap showing category co-occurrence."""
    plt.figure(figsize=(10, 8))
    
    # Create heatmap
    sns.heatmap(metrics['category_co_occurrence'], 
                annot=True, 
                fmt='d',
                cmap='YlOrRd',
                cbar_kws={'label': 'Co-occurrence Count'})
    
    plt.title('Innovation Category Co-occurrence Matrix', pad=20)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir['png'], 'category_co_occurrence.png'))
    plt.close()

def create_innovation_complexity_pie(metrics, output_dir):
    """Create a pie chart showing the distribution of innovation complexity."""
    plt.figure(figsize=(10, 8))
    
    dist = metrics['category_distribution']
    labels = [idx.replace('_', ' ').capitalize() for idx in dist.index]
    
    plt.pie(dist['percentage'], labels=labels, autopct='%1.1f%%',
            colors=sns.color_palette('pastel'))
    plt.title('Distribution of Innovation Complexity\n(Number of Categories per File)')
    
    plt.savefig(os.path.join(output_dir['png'], 'innovation_complexity.png'))
    plt.close()

def create_term_type_distribution(classifier, output_dir):
    """Create a stacked bar chart showing primary vs related term distribution."""
    stats = classifier.get_category_statistics()
    
    plt.figure(figsize=(12, 6))
    
    # Prepare data for stacked bar chart
    primary = stats['primary_term_matches']
    related = stats['related_term_matches']
    
    # Create stacked bar chart
    ax = plt.gca()
    ax.bar(stats['category'], primary, label='Primary Terms')
    ax.bar(stats['category'], related, bottom=primary, label='Related Terms')
    
    plt.title('Distribution of Primary vs Related Terms by Category', pad=20)
    plt.xlabel('Innovation Category')
    plt.ylabel('Number of Matches')
    plt.xticks(rotation=45, ha='right')
    plt.legend()
    
    # Add value labels
    for i in range(len(stats)):
        # Label for primary terms
        if primary[i] > 0:
            plt.text(i, primary[i]/2, str(primary[i]), ha='center', va='center')
        # Label for related terms
        if related[i] > 0:
            plt.text(i, primary[i] + related[i]/2, str(related[i]), ha='center', va='center')
    
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir['png'], 'term_type_distribution.png'))
    plt.close()

def create_visualizations(metrics, classifier, output_dir):
    """Create and save all visualization charts."""
    print("Generating visualization charts...")
    
    # Set the style for all plots
    sns.set_style("whitegrid")
    plt.rcParams['figure.facecolor'] = 'white'
    
    # Create individual charts
    create_category_distribution_chart(metrics, output_dir)
    create_co_occurrence_heatmap(metrics, output_dir)
    create_innovation_complexity_pie(metrics, output_dir)
    create_term_type_distribution(classifier, output_dir)

def save_classification_results(classifier, output_dir):
    """Save detailed classification results to CSV files."""
    # Save category statistics
    stats_df = classifier.get_category_statistics()
    stats_df.to_csv(os.path.join(output_dir['csv'], "detailed_category_statistics.csv"))
    
    # Save innovation patterns
    patterns_df = classifier.get_innovation_patterns()
    if not patterns_df.empty:
        patterns_df.to_csv(os.path.join(output_dir['csv'], "innovation_patterns.csv"), index=False)

def generate_innovation_metrics(results_df, total_files):
    """Generate detailed innovation metrics from the results."""
    # Files with any innovation
    files_with_innovation = len(results_df['rel_path'].unique())
    innovation_percentage = round((files_with_innovation / total_files * 100), 2)
    
    # Trial ID information
    trial_ids = results_df['trial_id'].unique()
    files_by_trial = results_df.groupby('trial_id')['rel_path'].nunique()
    
    # Category-level statistics
    category_stats = results_df.groupby('category').agg({
        'rel_path': 'nunique',
        'matched_term': 'count'
    }).rename(columns={
        'rel_path': 'unique_files',
        'matched_term': 'total_matches'
    })
    
    # Calculate percentages
    category_stats['percentage_of_files'] = round((category_stats['unique_files'] / total_files * 100), 2)
    category_stats['matches_per_file'] = round((category_stats['total_matches'] / category_stats['unique_files']), 2)
    
    # Calculate co-occurrence matrix
    file_category_matrix = pd.crosstab(results_df['rel_path'], results_df['category'])
    co_occurrence = pd.crosstab(results_df['rel_path'], results_df['category']).gt(0).astype(int)
    category_co_occurrence = co_occurrence.T.dot(co_occurrence)
    
    # Distribution of innovation categories per file
    innovation_counts = co_occurrence.sum(axis=1)
    category_distribution = pd.Series({
        '1_category': (innovation_counts == 1).sum(),
        '2_categories': (innovation_counts == 2).sum(),
        '3_or_more_categories': (innovation_counts >= 3).sum()
    })
    category_distribution_pct = round((category_distribution / files_with_innovation * 100), 2)
    
    return {
        'overall_metrics': {
            'total_files': total_files,
            'files_with_innovation': files_with_innovation,
            'innovation_percentage': innovation_percentage,
            'avg_categories_per_file': round(innovation_counts.mean(), 2),
            'unique_trial_ids': len(trial_ids)
        },
        'category_stats': category_stats,
        'category_co_occurrence': category_co_occurrence,
        'category_distribution': pd.DataFrame({
            'count': category_distribution,
            'percentage': category_distribution_pct
        }),
        'trial_stats': files_by_trial
    }

def save_innovation_metrics(metrics, output_dir):
    """Save innovation metrics to separate CSV files."""
    # Save overall metrics
    pd.DataFrame([metrics['overall_metrics']]).to_csv(
        os.path.join(output_dir['csv'], "overall_metrics.csv")
    )
    
    # Save category statistics
    metrics['category_stats'].to_csv(
        os.path.join(output_dir['csv'], "category_statistics.csv")
    )
    
    # Save co-occurrence matrix
    metrics['category_co_occurrence'].to_csv(
        os.path.join(output_dir['csv'], "category_co_occurrence.csv")
    )
    
    # Save category distribution
    metrics['category_distribution'].to_csv(
        os.path.join(output_dir['csv'], "category_distribution.csv")
    )
    
    # Save trial ID statistics
    metrics['trial_stats'].to_csv(
        os.path.join(output_dir['csv'], "trial_id_statistics.csv")
    )

def create_analysis_word_document(metrics, classifier, output_dir, timestamp):
    """Create a detailed Word document summarizing the analysis results."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Clinical Trial Innovation Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add timestamp
    doc.add_paragraph(f'Analysis Date: {timestamp}')
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    overall = metrics['overall_metrics']
    summary = doc.add_paragraph()
    summary.add_run('Analysis Overview:\n').bold = True
    summary.add_run(f"• {overall['files_with_innovation']} out of {overall['total_files']} files ({overall['innovation_percentage']}%) contain innovative techniques\n")
    summary.add_run(f"• Average innovation categories per file: {overall['avg_categories_per_file']}\n")
    summary.add_run(f"• Analysis includes {len(metrics['category_stats'])} innovation categories\n")
    summary.add_run(f"• Unique clinical trial IDs in analysis: {overall['unique_trial_ids']}\n")
    summary.add_run(f"• Files were processed recursively from all subdirectories\n")
    doc.add_paragraph()
    
    # Trial ID Summary
    if 'trial_stats' in metrics and not metrics['trial_stats'].empty:
        doc.add_heading('Clinical Trial ID Summary', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ['Trial ID', 'Number of Files']
        for i, header in enumerate(headers):
            header_cells[i].text = header
        
        for trial_id, file_count in metrics['trial_stats'].items():
            if pd.notnull(trial_id): # Skip None/NaN trial IDs
                row_cells = table.add_row().cells
                row_cells[0].text = str(trial_id)
                row_cells[1].text = str(file_count)
        
        doc.add_paragraph()
    
    # Innovation Categories Analysis
    doc.add_heading('Innovation Categories Analysis', level=1)
    
    # Category Distribution
    doc.add_heading('Category Distribution', level=2)
    img_path = os.path.join(output_dir['png'], 'category_distribution.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 1: Distribution of Innovation Categories Across Files')
    
    # Add category statistics table
    doc.add_heading('Category Statistics', level=2)
    stats = metrics['category_stats']
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    headers = ['Category', 'Files', '% of Total', 'Matches/File']
    for i, header in enumerate(headers):
        header_cells[i].text = header
    
    for category, row in stats.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = str(row['unique_files'])
        row_cells[2].text = f"{row['percentage_of_files']}%"
        row_cells[3].text = str(row['matches_per_file'])
    
    doc.add_paragraph()
    
    # Co-occurrence Analysis
    doc.add_heading('Innovation Co-occurrence Analysis', level=2)
    img_path = os.path.join(output_dir['png'], 'category_co_occurrence.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 2: Innovation Category Co-occurrence Matrix')
    
    # Innovation Complexity
    doc.add_heading('Innovation Complexity Analysis', level=2)
    img_path = os.path.join(output_dir['png'], 'innovation_complexity.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 3: Distribution of Innovation Complexity')
    
    # Term Type Distribution
    doc.add_heading('Term Type Distribution', level=2)
    img_path = os.path.join(output_dir['png'], 'term_type_distribution.png')
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 4: Distribution of Primary vs Related Terms by Category')
    
    # Detailed Category Analysis
    doc.add_heading('Detailed Category Analysis', level=1)
    stats = classifier.get_category_statistics()
    for _, row in stats.iterrows():
        doc.add_heading(row['category'], level=2)
        p = doc.add_paragraph()
        p.add_run('Primary Term Matches: ').bold = True
        p.add_run(f"{row['primary_term_matches']}\n")
        p.add_run('Related Term Matches: ').bold = True
        p.add_run(f"{row['related_term_matches']}\n")
        p.add_run('Most Common Terms: ').bold = True
        common_terms = list(row['most_common_terms'].keys())[:5]
        p.add_run(', '.join(common_terms))
        doc.add_paragraph()
    
    # Generated Files
    doc.add_heading('Generated Analysis Files', level=1)
    files_section = doc.add_paragraph()
    
    # List files by directory
    for dir_type, dir_path in output_dir.items():
        if dir_type != 'report':  # Skip report directory in listing
            files_section.add_run(f'\n{dir_type.upper()} Files:\n').bold = True
            files = [f for f in os.listdir(dir_path) if os.path.isfile(os.path.join(dir_path, f))]
            for file in sorted(files):
                files_section.add_run(f"• {file}\n")
    
    # Save the document
    doc_path = os.path.join(output_dir['report'], 'innovation_analysis_report.docx')
    doc.save(doc_path)
    return doc_path

def process_single_pdf(args):
    """Process a single PDF file and return the results."""
    pdf_path, rel_path, filename, lexicon_terms, threshold = args
    
    # Extract trial ID from the path (e.g., NCT04261244)
    trial_id = None
    path_parts = pdf_path.split(os.sep)
    for part in path_parts:
        if part.startswith("NCT"):
            trial_id = part
            break
    
    file_start_time = time.time()
    results = []
    
    # Count terms and track context
    matches = process_pdf_for_terms(pdf_path, lexicon_terms, threshold)
    
    # Only process and add results if there are matches
    if matches:
        # Add filename, relative path, and trial ID to each match
        for match in matches:
            match["filename"] = filename
            match["rel_path"] = rel_path
            match["trial_id"] = trial_id
        results = matches
        
        file_duration = time.time() - file_start_time
        return {
            "matches": results,
            "rel_path": rel_path,
            "trial_id": trial_id,
            "has_matches": True,
            "match_count": len(matches),
            "duration": file_duration
        }
    else:
        file_duration = time.time() - file_start_time
        return {
            "matches": [],
            "rel_path": rel_path,
            "trial_id": trial_id,
            "has_matches": False,
            "match_count": 0,
            "duration": file_duration
        }

def process_all_pdfs(pdf_folder, lexicon_file, output_folder, threshold=85, workers=None):
    """Process all PDFs for term counting, context, and highlighting."""
    start_time = time.time()
    
    # Set default number of workers if not specified
    if workers is None:
        workers = max(1, multiprocessing.cpu_count() - 1)
        
    print(f"Using {workers} worker processes for parallel processing")
    
    # Create timestamped subdirectory
    timestamp = pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')
    timestamped_output_dir = os.path.join(output_folder, timestamp)
    os.makedirs(timestamped_output_dir, exist_ok=True)
    
    # Create organized subdirectories
    output_dirs = create_output_directories(timestamped_output_dir)

    print(f"\nLoading lexicon from {lexicon_file}...")
    lexicon_load_start = time.time()
    lexicon_terms = load_lexicon(lexicon_file)
    print(f"Loaded {len(lexicon_terms)} lexicon terms in {format_time(time.time() - lexicon_load_start)}")
    
    all_results = []
    
    # Walk through all subdirectories to find PDF files
    pdf_files = []
    for root, _, files in os.walk(pdf_folder):
        for filename in files:
            if filename.endswith(".pdf"):
                # Store the full path and the relative path for later use
                full_path = os.path.join(root, filename)
                rel_path = os.path.relpath(full_path, pdf_folder)
                pdf_files.append((full_path, rel_path, filename))
    
    total_files = len(pdf_files)
    
    print(f"\nProcessing {total_files} PDF files from {pdf_folder} and its subdirectories...")
    
    # Create processing tasks
    processing_tasks = [(pdf_path, rel_path, filename, lexicon_terms, threshold) 
                        for pdf_path, rel_path, filename in pdf_files]
    
    # Initialize progress bar
    progress_bar = tqdm(total=total_files, desc="Processing PDFs", unit="file")
    
    # Use ProcessPoolExecutor for parallel processing
    with ProcessPoolExecutor(max_workers=workers) as executor:
        futures = {executor.submit(process_single_pdf, task): task for task in processing_tasks}
        
        for future in as_completed(futures):
            result = future.result()
            rel_path = result["rel_path"]
            
            # Update progress bar description with current file and elapsed time
            elapsed = time.time() - start_time
            progress_bar.set_description(f"Processed {rel_path} (Elapsed: {format_time(elapsed)})")
            
            if result["has_matches"]:
                all_results.extend(result["matches"])
                
                # Extract data needed for highlighting
                pdf_path = futures[future][0]
                filename = futures[future][2]
                matches = result["matches"]
                trial_id = result["trial_id"]
                
                # Create highlighted PDF
                rel_dir = os.path.dirname(rel_path)
                if rel_dir:
                    highlighted_pdf_dir = os.path.join(output_dirs['pdf'], rel_dir)
                    os.makedirs(highlighted_pdf_dir, exist_ok=True)
                    highlighted_pdf_path = os.path.join(highlighted_pdf_dir, f"highlighted_{filename}")
                else:
                    # If trial ID was extracted, create a directory for it
                    if trial_id:
                        highlighted_pdf_dir = os.path.join(output_dirs['pdf'], trial_id)
                        os.makedirs(highlighted_pdf_dir, exist_ok=True)
                        highlighted_pdf_path = os.path.join(highlighted_pdf_dir, f"highlighted_{filename}")
                    else:
                        highlighted_pdf_path = os.path.join(output_dirs['pdf'], f"highlighted_{filename}")
                    
                highlight_terms_in_pdf(pdf_path, matches, highlighted_pdf_path)
                progress_bar.write(f"Created highlighted PDF with {result['match_count']} matches")
            else:
                progress_bar.write(f"No matches found in {rel_path}")
            
            progress_bar.write(f"File processing time: {format_time(result['duration'])}")
            
            # Update progress bar
            progress_bar.update(1)
    
    # Close progress bar
    progress_bar.close()
    
    # Only create output files if there are any results
    if all_results:
        print("\nGenerating reports...")
        report_start_time = time.time()
        
        # Create classifier instance
        classifier = MatchClassifier(all_results)
        
        # Save detailed results to CSV
        results_df = pd.DataFrame(all_results)
        csv_output_path = os.path.join(output_dirs['csv'], "term_locations_with_context.csv")
        results_df.to_csv(csv_output_path, index=False)
        print(f"Term location results saved to: {csv_output_path}")

        # Generate and save innovation metrics
        print("Generating innovation metrics...")
        metrics = generate_innovation_metrics(results_df, total_files)
        save_innovation_metrics(metrics, output_dirs)
        
        # Save classification results
        save_classification_results(classifier, output_dirs)
        
        # Create visualizations
        create_visualizations(metrics, classifier, output_dirs)
        
        # Generate Word document report
        print("\nGenerating Word document report...")
        doc_path = create_analysis_word_document(metrics, classifier, output_dirs, timestamp)
        print(f"Analysis report saved to: {doc_path}")
        
        # Print summary of findings
        overall = metrics['overall_metrics']
        print(f"\nInnovation Analysis Summary:")
        print(f"- {overall['files_with_innovation']} out of {overall['total_files']} files ({overall['innovation_percentage']}%) contain innovative techniques")
        print(f"- Average number of innovation categories per file: {overall['avg_categories_per_file']}")
        print(f"- Files were processed recursively from {pdf_folder} and all its subdirectories")
        print(f"- Unique clinical trial IDs identified: {overall['unique_trial_ids']}")
        
        # Distribution summary
        dist = metrics['category_distribution']
        print("\nDistribution of innovation categories:")
        for idx, row in dist.iterrows():
            category_name = idx.replace('_', ' ').capitalize()
            print(f"- {category_name}: {row['count']} files ({row['percentage']}%)")
        
        # Trial ID summary
        if 'trial_stats' in metrics and not metrics['trial_stats'].empty:
            print("\nTop clinical trials by file count:")
            top_trials = metrics['trial_stats'].sort_values(ascending=False).head(5)
            for trial_id, count in top_trials.items():
                if pd.notnull(trial_id):  # Skip None/NaN trial IDs
                    print(f"- {trial_id}: {count} files")
        
        # Term type distribution summary
        stats = classifier.get_category_statistics()
        print("\nTerm type distribution:")
        for _, row in stats.iterrows():
            print(f"\n{row['category']}:")
            print(f"- Primary term matches: {row['primary_term_matches']}")
            print(f"- Related term matches: {row['related_term_matches']}")
            print(f"- Most common terms: {', '.join(list(row['most_common_terms'].keys())[:3])}")
        
        print(f"\nReport generation time: {format_time(time.time() - report_start_time)}")
        print("\nOutput files have been organized into the following directories:")
        print(f"- Visualizations: {output_dirs['png']}")
        print(f"- Data files: {output_dirs['csv']}")
        print(f"- Highlighted PDFs: {output_dirs['pdf']}")
        print(f"- Analysis report: {output_dirs['report']}")
    else:
        print("\nNo matches found in any files")
    
    total_duration = time.time() - start_time
    print(f"\nTotal processing time: {format_time(total_duration)}")
    print(f"Average time per file: {format_time(total_duration/total_files)}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Process PDFs for term counting, context, and highlighting")
    parser.add_argument("pdf_folder", help="Folder containing PDF files to process")
    parser.add_argument("lexicon_file", help="CSV file containing lexicon terms to match")
    parser.add_argument("output_folder", help="Folder for outputting results")
    parser.add_argument("--threshold", type=int, default=85, help="Matching threshold (default: 85)")
    parser.add_argument("--workers", type=int, help="Number of worker processes for parallel processing")
    
    args = parser.parse_args()
    
    process_all_pdfs(args.pdf_folder, args.lexicon_file, args.output_folder, args.threshold, args.workers)